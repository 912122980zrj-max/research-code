# -*- coding: utf-8 -*-
"""
Build MANUSCRIPT_FINAL_v5_ZH.docx / MANUSCRIPT_FINAL_v5_EN.docx from the
Markdown sources using the bundled python-docx runtime.

Design preset: narrative_proposal (Calibri 11 pt, justified, 1.333 line
spacing, blue heading ladder), with named manuscript overrides:
  - manuscript_title   18 pt bold #1F3A5F, centered
  - figure_placeholder shaded paragraph (#F2F4F7) + left blue bar
  - table_caption      10 pt bold, table_note 9 pt italic gray
  - zh_body_font       eastAsia = SimSun (宋体) for the Chinese draft
Real Word numbering is used for numbered limitations and bullet lists.
"""

import re
import sys
import io
import os
import zipfile
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

ROOT = Path(r"E:\sheng xin\tbi")
MD_ZH = ROOT / "reports" / "manuscript" / "MANUSCRIPT_FINAL_v7_ZH.md"
MD_EN = ROOT / "reports" / "manuscript" / "MANUSCRIPT_FINAL_v7_EN.md"
OUT_ZH = ROOT / "reports" / "manuscript" / "MANUSCRIPT_FINAL_v7_ZH.docx"
OUT_EN = ROOT / "reports" / "manuscript" / "MANUSCRIPT_FINAL_v7_EN.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TITLE_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Markdown -> blocks
# ---------------------------------------------------------------------------

def parse_md(text):
    lines = text.splitlines()
    blocks = []
    i = 0
    in_refs = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            blocks.append(("table", rows))
            continue
        if line.startswith("# "):
            blocks.append(("title", line[2:].strip()))
        elif line.startswith("## "):
            h = line[3:].strip()
            in_refs = h in ("References", "参考文献")
            blocks.append(("h2", h))
        elif line.startswith("### "):
            t = line[4:].strip()
            in_refs = False
            if t.startswith(("表 ", "Table ", "补充表", "Supplementary Table")):
                blocks.append(("caption", t))
            else:
                blocks.append(("h3", t))
        elif line == "---":
            pass
        elif line.startswith("> "):
            blocks.append(("note", line[2:].strip()))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        elif re.match(r"^\d+\.\s", line):
            if in_refs:
                blocks.append(("para", line))
            else:
                blocks.append(("numbered", re.sub(r"^\d+\.\s", "", line)))
        elif "插入位置" in line or "insertion:" in line:
            blocks.append(("placeholder", line))
        elif len(blocks) < 6 and line.startswith("**") and len(line) < 80 and (
                "：" in line or (re.match(r"^\*\*[^*]+:\*\*", line))):
            blocks.append(("meta", line))
        elif line.startswith(("图 ", "**图 ", "**Figure", "Figure ")) and len(line) > 8:
            blocks.append(("legend", line))
        elif line.startswith(("表 ", "Table ")) and len(line) < 100:
            blocks.append(("caption", line))
        elif line.startswith("*") and not line.startswith("**"):
            blocks.append(("table_note", line.lstrip("*").strip()))
        elif line.startswith("**软件引用") or line.startswith("**Software"):
            blocks.append(("para", line))
        elif line.startswith("**") and ("**：" in line or re.match(r"^\*\*[^*]+：", line)):
            blocks.append(("para", line))
        else:
            blocks.append(("para", line))
        i += 1
    return blocks


def strip_md(s):
    return s.replace("**", "").replace("`", "")


def add_runs(p, text, base_bold=False, size=None, color=None, italic=False,
             font=None):
    # minimal **bold** inline parser
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = base_bold
        txt = part
        if part.startswith("**") and part.endswith("**"):
            bold = True
            txt = part[2:-2]
        txt = txt.replace("`", "")
        if not txt:
            continue
        r = p.add_run(txt)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        if font:
            r.font.name = font
            r._element.rPr.rFonts.set(qn("w:ascii"), font)
            r._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    return p


def split_legend_lead(text):
    t = strip_md(text)
    if t.startswith("图 "):
        idx = t.find("。")
        if idx > 0:
            return t[: idx + 1], t[idx + 1 :].strip()
    m = re.match(r"^(Figure \d+\. [^.]+\.)", t)
    if m:
        return m.group(1), t[m.end() :].strip()
    return t, ""


def parse_table(rows):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if all(re.match(r"^:?-{2,}:?$", c) for c in cells):
            continue
        parsed.append(cells)
    return parsed


# ---------------------------------------------------------------------------
# Numbering part
# ---------------------------------------------------------------------------

NUMBERING_XML = """<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="100">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="101">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="\u2022"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="100"/></w:num>
  <w:num w:numId="2"><w:abstractNumId w:val="101"/></w:num>
</w:numbering>"""


def set_num(p, num_id):
    pPr = p._p.get_or_add_pPr()
    numPr = parse_xml(
        '<w:numPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:ilvl w:val="0"/><w:numId w:val="%d"/></w:numPr>' % num_id
    )
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# Table geometry
# ---------------------------------------------------------------------------

def table_widths(ncols):
    weights = {
        2: [0.5, 0.5],
        3: [0.08, 0.42, 0.50],
        5: [0.12, 0.20, 0.24, 0.22, 0.22],
        6: [0.13, 0.14, 0.15, 0.11, 0.19, 0.28],
    }.get(ncols, [1.0 / ncols] * ncols)
    widths = [int(round(w * 9360)) for w in weights]
    widths[-1] = 9360 - sum(widths[:-1])
    return widths


def set_table_geometry(table, widths):
    tbl = table._tbl
    tblPr = tbl.tblPr
    # width + indent
    for tag, val in (("w:tblW", "9360"), ("w:tblInd", "120")):
        el = tblPr.find(qn(tag))
        if el is None:
            el = parse_xml(
                '<w:%s xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="dxa" w:w="%s"/>'
                % (tag.split(":")[1], val)
            )
            tblPr.append(el)
        else:
            el.set(qn("w:w"), val)
            el.set(qn("w:type"), "dxa")
    # grid
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = parse_xml(
        '<w:tblGrid xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        + "".join('<w:gridCol w:w="%d"/>' % w for w in widths)
        + "</w:tblGrid>"
    )
    tbl.insert(list(tbl).index(tblPr) + 1, grid)
    # cell widths + margins + vertical alignment
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = parse_xml(
                    '<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="dxa"/>'
                )
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[idx]))
            tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tcPr2 = tcPr.find(qn("w:tcMar"))
            if tcPr2 is None:
                tcPr2 = parse_xml(
                    '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<w:top w:w="80" w:type="dxa"/><w:start w:w="120" w:type="dxa"/>'
                    '<w:bottom w:w="80" w:type="dxa"/><w:end w:w="120" w:type="dxa"/></w:tcMar>'
                )
                tcPr.append(tcPr2)


def shade_cell(cell, fill="F4F6F9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="%s"/>'
        % fill
    )
    tcPr.append(shd)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(
        '<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    ))


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def set_style_font(style, font="Calibri", east="宋体"):
    style.font.name = font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(
            '<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), east)


def build_doc(blocks, out_path, zh):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(sec, attr, Inches(1.0))
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, east="宋体" if zh else "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        st = doc.styles[name]
        set_style_font(st, east="黑体" if zh else "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    first = True
    for kind, payload in blocks:
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.2
            add_runs(p, strip_md(payload), base_bold=True, size=18,
                     color=TITLE_BLUE)
        elif kind == "meta":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_runs(p, payload, size=10, color=GRAY)
        elif kind == "h2":
            doc.add_heading(strip_md(payload), level=1)
        elif kind == "h3":
            doc.add_heading(strip_md(payload), level=2)
        elif kind == "note":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(8)
            add_runs(p, strip_md(payload), size=10, color=GRAY, italic=True)
        elif kind == "bullet":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            set_num(p, 2)
            add_runs(p, payload)
        elif kind == "numbered":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            set_num(p, 1)
            add_runs(p, payload)
        elif kind == "placeholder":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.line_spacing = 1.15
            shd = parse_xml(
                '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'w:val="clear" w:color="auto" w:fill="F2F4F7"/>'
            )
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:left w:val="single" w:sz="18" w:space="4" w:color="2E74B5"/></w:pBdr>'
            )
            pPr.append(shd)
            pPr.append(pBdr)
            add_runs(p, payload, size=10)
        elif kind == "legend":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            lead, rest = split_legend_lead(payload)
            add_runs(p, lead, base_bold=True)
            if rest:
                add_runs(p, " " + rest)
        elif kind == "caption":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            add_runs(p, strip_md(payload), base_bold=True, size=10)
        elif kind == "table_note":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            add_runs(p, payload, size=9, color=GRAY, italic=True)
        elif kind == "para":
            p = doc.add_paragraph()
            add_runs(p, payload)
        elif kind == "table":
            parsed = parse_table(payload)
            if not parsed:
                continue
            ncols = len(parsed[0])
            widths = table_widths(ncols)
            table = doc.add_table(rows=len(parsed), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(parsed):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.paragraph_format.space_after = Pt(2)
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.line_spacing = 1.1
                    is_num = bool(re.match(
                        r"^[0-9+\-−.,%×⁻⁹()/·\s]+$", cell_text))
                    para.alignment = (WD_ALIGN_PARAGRAPH.CENTER if is_num
                                      else WD_ALIGN_PARAGRAPH.LEFT)
                    add_runs(para, cell_text, base_bold=(ri == 0), size=9.5)
                if ri == 0:
                    repeat_header(table.rows[0])
                    for c in table.rows[0].cells:
                        shade_cell(c)
            set_table_geometry(table, widths)
        else:
            continue

    # page number footer (quiet, right aligned)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run()
    fld = parse_xml(
        '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
    )
    fp._p.append(fld)
    for r in fp.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY

    doc.save(str(out_path))
    add_numbering_part_zip(out_path)
    print("saved", out_path)


def add_numbering_part_zip(docx_path):
    """Append custom list definitions into the template's word/numbering.xml."""
    tmp = docx_path.with_suffix(".tmp.docx")
    payload = NUMBERING_XML.encode("utf-8")
    with zipfile.ZipFile(str(docx_path)) as zin, \
            zipfile.ZipFile(str(tmp), "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/numbering.xml":
                s = data.decode("utf-8")
                if '<w:num w:numId="1">' not in s:
                    fragment = payload.decode("utf-8")
                    fragment = fragment[fragment.find("<w:abstractNum"):]
                    fragment = fragment.replace("</w:numbering>", "")
                    s = s.replace("</w:numbering>", fragment + "</w:numbering>")
                data = s.encode("utf-8")
            zout.writestr(item, data)
    os.replace(str(tmp), str(docx_path))


def audit(doc_path):
    doc = Document(str(doc_path))
    n_para = len(doc.paragraphs)
    n_tbl = len(doc.tables)
    issues = []
    for t in doc.tables:
        tbl = t._tbl
        grid = tbl.find(qn("w:tblGrid"))
        if grid is None:
            issues.append("table without tblGrid")
            continue
        widths = [int(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol"))]
        if sum(widths) != 9360:
            issues.append("grid sum %d" % sum(widths))
        for row in t.rows:
            for idx, cell in enumerate(row.cells):
                tcPr = cell._tc.tcPr
                tcW = tcPr.find(qn("w:tcW")) if tcPr is not None else None
                if tcW is None or int(tcW.get(qn("w:w"))) != widths[idx]:
                    issues.append("cell width mismatch row %d col %d" % (0, idx))
                    break
    text = "\n".join(p.text for p in doc.paragraphs)
    zh = "ZH" in doc_path.name
    markers = ("【图 1 插入位置", "【图 2 插入位置", "【图 3 插入位置",
               "【图 4 插入位置") if zh else (
               "【Figure 1 insertion", "【Figure 2 insertion",
               "【Figure 3 insertion", "【Figure 4 insertion",
               )
    for marker in markers:
        if marker not in text:
            issues.append("missing marker " + marker)
    print("audit", doc_path.name, "paras", n_para, "tables", n_tbl,
          "issues", issues or "none")


if __name__ == "__main__":
    if "--debug-parse" in sys.argv:
        for name, path in (("ZH", MD_ZH), ("EN", MD_EN)):
            blocks = parse_md(path.read_text(encoding="utf-8"))
            kinds = [k for k, _ in blocks]
            print(name, "blocks", len(blocks),
                  "placeholder", kinds.count("placeholder"),
                  "legend", kinds.count("legend"),
                  "table", kinds.count("table"))
            for k, v in blocks:
                if k == "placeholder":
                    print("  ", v[:50])
        sys.exit(0)
    zh_blocks = parse_md(MD_ZH.read_text(encoding="utf-8"))
    en_blocks = parse_md(MD_EN.read_text(encoding="utf-8"))
    build_doc(zh_blocks, OUT_ZH, zh=True)
    build_doc(en_blocks, OUT_EN, zh=False)
    audit(OUT_ZH)
    audit(OUT_EN)
