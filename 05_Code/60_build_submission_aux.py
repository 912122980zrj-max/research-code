# -*- coding: utf-8 -*-
# ============================================================
# 60_build_submission_aux.py —— legacy (v4) auxiliary builder
# 注意：v6 补充表工作簿由 68_build_supplementary_v6.py 生成；
# 本脚本生成的 Supplementary_Tables_S1-S12.xlsx 已不再用于投稿包
# （仅保留 cover letter DOCX 生成逻辑；如需重跑请先确认 v6 编号）。
# ============================================================
"""Build submission auxiliary files:
1) S11 hub-gene DEG table from the v4 deg_all.csv
2) Supplementary_Tables_S1-S12.xlsx from raw CSVs
3) COVER_LETTER_EN.docx
"""
import sys, io, csv
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

ROOT = Path(r"E:\sheng xin\tbi")
RAW = ROOT / "submission" / "03_Supplementary" / "Tables" / "Raw_CSV"
OUT_XLSX = ROOT / "submission" / "03_Supplementary" / "Tables" / "Supplementary_Tables_S1-S12.xlsx"
COVER = ROOT / "submission" / "01_Manuscript" / "COVER_LETTER_EN.docx"


def read_csv(p):
    with open(p, encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


# ---- 1) S11 hub genes ----
deg = read_csv(ROOT / "results" / "03_DEG_Analysis" / "deg_all.csv")
hub = ["Trp53", "Egfr", "Akt1", "Actb", "Tnf"]
s11_path = RAW / "S11_hub_genes_deg.csv"
with open(s11_path, "w", newline="", encoding="utf-8") as f:
    w = csv.writer(f)
    w.writerow(["Mouse gene", "Human symbol", "log2FC", "P", "BH adjP",
                "DEG (|log2FC|>0.585, adjP<0.05)", "Note"])
    for g in hub:
        row = next((r for r in deg if r["gene"] == g), None)
        if row is None:
            w.writerow([g, {"Trp53": "TP53", "Egfr": "EGFR", "Akt1": "AKT1",
                            "Actb": "ACTB", "Tnf": "TNF"}[g],
                        "NA", "NA", "NA", "No",
                        "Absent from the coverage-filtered matrix (not evaluable)"])
        else:
            w.writerow([g, {"Trp53": "TP53", "Egfr": "EGFR", "Akt1": "AKT1",
                            "Actb": "ACTB", "Tnf": "TNF"}[g],
                        round(float(row["logFC"]), 3),
                        format(float(row["P.Value"]), ".3g"),
                        format(float(row["adj.P.Val"]), ".3g"),
                        "No", ""])
print("S11 written:", s11_path)

# ---- 2) XLSX ----
wb = Workbook()
wb.remove(wb.active)
header_fill = PatternFill("solid", fgColor="DCE6F1")
header_font = Font(bold=True)

sheet_map = [
    ("S1_CohortDEG_Overlap", ["S1_cohort_deg_overlap.csv"]),
    ("S2_GSE41345_DEG", ["S2_GSE41345_empty_deg.csv"]),
    ("S3_GO_BP", ["S3_GO_BP_keep.csv"]),
    ("S3_AhR_Fisher", ["S3_ahr_fisher.csv"]),
    ("S4_ENet_Coefficients", ["S4_enet_coefficients.csv"]),
    ("S4_Frozen_Coefficients", ["S4_frozen_coefficients.csv"]),
    ("S5_LOCO", ["S5_loco_auc.csv", "S5_single_cohort_model_auc.csv"]),
    ("S6_Deconvolution_Corr", ["S6_deconvolution_method_corr.csv"]),
    ("S7_Composition", ["S7_composition_stats.csv", "S7_composition_stats_d7.csv"]),
    ("S8_Docking", ["S8_docking_energies.csv", "S8_docking_box_scores.csv",
                    "S8_redocking_rmsd.csv", "S8_decoy_scores.csv",
                    "S8_decoy_summary.csv", "S8_blind_control.csv"]),
    ("S9_Astro", ["S9_astro_subcluster_stats.csv", "S9_astro_annotation.csv",
                  "S9_astro_purity.csv"]),
    ("S11_HubGenes", ["S11_hub_genes_deg.csv"]),
    ("S12_Algorithm_Benchmark", ["S12_algorithm_benchmark.csv"]),
]

for sheet_name, files in sheet_map:
    ws = wb.create_sheet(sheet_name[:31])
    first = True
    for fn in files:
        rows = read_csv(RAW / fn)
        if not rows:
            continue
        cols = list(rows[0].keys())
        if first:
            ws.append(cols)
            for c in range(1, len(cols) + 1):
                cell = ws.cell(row=1, column=c)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(vertical="center")
            first = False
        else:
            ws.append([""])
            ws.append([f"--- {fn} ---"])
        for r in rows:
            ws.append([r.get(c, "") for c in cols])
    for c in range(1, ws.max_column + 1):
        letter = get_column_letter(c)
        width = max((len(str(ws.cell(row=r, column=c).value or ""))
                     for r in range(1, min(ws.max_row, 200) + 1)), default=8)
        ws.column_dimensions[letter].width = min(max(width + 2, 8), 45)
    ws.freeze_panes = "A2"

wb.save(str(OUT_XLSX))
print("XLSX written:", OUT_XLSX)

# ---- 3) Cover letter DOCX ----
from docx import Document
from docx.shared import Pt, Inches

doc = Document()
sec = doc.sections[0]
for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
    setattr(sec, attr, Inches(1.0))
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = 1
run = title.add_run("Dear Editor,")
run.bold = True
run.font.size = Pt(12)

body = [
    "We would like to submit the enclosed manuscript, \u201cPredicted benzo[a]pyrene "
    "targets show modest overlap above measurability-matched background with traumatic "
    "brain injury transcriptomes: "
    "a multi-cohort transcriptomic stress test\u201d, for consideration as an original "
    "research article.",
    "Network toxicology commonly interprets the intersection between a compound\u2019s "
    "computationally predicted targets and disease differentially expressed genes as "
    "evidence of an exposure\u2013outcome molecular bridge. Using six public transcriptomic "
    "datasets (four mouse bulk cohorts, one human aging cohort, and one human snRNA-seq "
    "dataset), we quantified this intersection for benzo[a]pyrene (BaP) and acute "
    "traumatic brain injury (TBI) under a coverage-filtered, batch-corrected pipeline. "
    "traumatic brain injury (TBI) under an orthology-mapped, coverage-filtered, "
    "batch-corrected pipeline with a prespecified permutation background test. The "
    "measured overlap was modest but statistically above measurability-matched "
    "background (7 weak-evidence candidate genes; 26.1-fold enrichment; empirical "
    "P<0.0001), and the strongest discriminative features (P2RY6/HSD11B1) belonged to "
    "endogenous injury programs rather than predicted exposure targets. Relative to "
    "earlier versions of this work, we corrected the GSE41345 strain/time-point "
    "description (ICR, 14 d), replaced case-insensitive symbol matching with Ensembl "
    "Compara 1:1 orthology (the previously reported \u201call tier A unmeasurable\u201d "
    "was a symbol-matching artifact; HPRT1/Hprt is measurable but not a DEG), removed "
    "the non-evaluable AhR Fisher P value, and added the permutation and sensitivity "
    "analyses. Because no cohort contained individual exposure measurements, the study "
    "is explicitly hypothesis-generating and is framed as a methodological stress test "
    "of a widely used analysis strategy.",
    "We believe this work will interest readers concerned with the reproducibility of "
    "network-toxicology and transcriptomic biomarker studies, and with the interpretation "
    "limits of computational target predictions in acute injury. All data are public GEO "
    "datasets; analysis code and processed data will be deposited in a public repository "
    "upon acceptance.",
    "The manuscript has not been published or submitted elsewhere, and all authors "
    "approved the submission. We declare no competing interests.",
    "Thank you for your consideration. We look forward to your response.",
]
for para in body:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(para)

sign = doc.add_paragraph()
sign.paragraph_format.space_before = Pt(12)
sign.add_run("Zuoren Jie\nCorresponding author\n[Affiliation] [Email]")
doc.save(str(COVER))
print("Cover letter written:", COVER)
